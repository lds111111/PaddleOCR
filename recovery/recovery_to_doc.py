import os
from copy import deepcopy

from collections import Counter

import docx
from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt

from ppstructure.recovery.table_process import HtmlToDocx

from ppocr.utils.logging import get_logger
logger = get_logger()


def convert_info_docx(args, img, res, save_folder, img_name, save_interval=100):
    doc = Document()
    doc.styles['Normal'].font.name = args.set_doc_english_font
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), args.set_doc_chinese_font)
    doc.styles['Normal'].font.size = shared.Pt(args.set_doc_base_font_size)


    flag = 1

    tmp_img_idx = 0

    excel_save_folder = os.path.join(save_folder, img_name)

    docx_path = os.path.join(save_folder, '{}_ocr.docx'.format(img_name))

    for i, region in enumerate(res):

        if region['type'] in ['header', 'footer']:
            continue

        img_idx = region['img_idx']

        if img_idx != tmp_img_idx:
            tmp_img_idx = img_idx
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)

        if (flag == 2 and region['layout'] == 'single') or (flag == 1 and region['layout'] == 'double'):
            add_section_layout(doc, region['layout'])
            flag = 1 if flag == 2 else 2

        region_type = region['type'].lower()

        if region_type == 'figure':
            img_path = os.path.join(excel_save_folder, '{}_{}.jpg'.format(region['bbox'], img_idx))
            add_image_to_doc(doc, img_path, flag)

        elif region_type == 'title':
            add_section_layout(doc, region["page_max_layout_flag"])
            doc.add_heading(region['res'][0]['text'])

        elif region_type == 'table':
            doc.styles['Normal'].font.size = shared.Pt(args.set_doc_table_font_size)
            parser = HtmlToDocx()
            parser.table_style = 'TableGrid'
            parser.handle_table(region['res']['html'], doc)

        else:
            if region_type in ['figure_caption', 'table_caption']:
                add_section_layout(doc, region["page_max_layout_flag"])

            for line in region['res']:
                paragraph = doc.add_paragraph()
                text_run = paragraph.add_run(line.get('text'))
                text_run.font.size = shared.Pt(args.set_doc_figure_and_table_caption_font_size if region_type in ['figure_caption', 'table_caption'] else args.set_doc_text_font_size)
                if region_type  in ['figure_caption', 'table_caption']:
                    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                else:
                    paragraph.paragraph_format.first_line_indent = text_run.font.size * 2
                    paragraph.paragraph_format.line_spacing = args.set_doc_text_line_spacing
                    #设置段落下边距
                    paragraph.paragraph_format.space_after = Pt(args.set_doc_text_space_after)


        # Save the document every save_interval regions
        if i % save_interval == 0:
            doc.save(docx_path)
            logger.info('docx save batch: {}'.format(i))
    # Save the document one last time at the end to make sure all regions are written
    doc.save(docx_path)
    logger.info('Save the document one last time')
    logger.info('docx save to {}'.format(docx_path))


#该函数将在从左到右，从上到下的顺序对文本框进行排序，并确定它们的布局（单列或双列）。
def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    #获取结果中的框的数量。
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]['layout'] = 'single'
        return res

    #对框进行排序，排序标准是每个框左上角点的 y 坐标和 x 坐标。
    sorted_boxes = sorted(res, key=lambda x: (x['bbox'][1], x['bbox'][0]))
    #将排序后的框复制到新的列表中。
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0
    # while 循环处理每个框的布局。根据每个框的位置，确定它是在单列布局中，还是在双列布局中的左边或右边。
    # 如果一个框不能被明确地分配到左列或右列，那么它就被分配到单列布局中。
    # 判断文本框是应该被分配到左列、右列或者单列中。这通过判断每个框的左上角和右下角的坐标相对于页面宽度的位置来实现。
    # 如果一个框的左上角和右下角的 x 坐标都小于页面宽度的 3/4，那么它被判定为左列；如果一个框的左上角的 x 坐标大于页面宽度的 1/4。
    # 并且其右下角的 x 坐标大于页面宽度的一半，那么它被判定为右列；否则，它被判定为单列。
    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if _boxes[i]['bbox'][1] > _boxes[i - 1]['bbox'][3] and _boxes[i][
                    'bbox'][0] < w / 2 and _boxes[i]['bbox'][2] > w / 2:
                new_res += res_left
                new_res += res_right
                _boxes[i]['layout'] = 'single'
                new_res.append(_boxes[i])
            else:
                if _boxes[i]['bbox'][2] > w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]['bbox'][0] < w / 2:
                    _boxes[i]['layout'] = 'double'
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []

            tmp_layout = []
            # 如果处理整页,返回字典数组中['layout']值出现最多的值
            for i in range(len(_boxes)):
                if _boxes[i]['type'] == 'text' or _boxes[i]['type'] == 'figure' or _boxes[i]['type'] == 'table':
                    tmp_layout.append(_boxes[i]['layout'])

            for i in range(len(_boxes)):
                _boxes[i]["page_max_layout_flag"] = Counter(tmp_layout).most_common(1)[0][0]

            break


        elif _boxes[i]['bbox'][0] < w / 4 and _boxes[i]['bbox'][2] < 3 * w / 4:
            _boxes[i]['layout'] = 'double'
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]['bbox'][0] > w / 4 and _boxes[i]['bbox'][2] > w / 2:
            _boxes[i]['layout'] = 'double'
            res_right.append(_boxes[i])
            i += 1
        else:
            new_res += res_left
            new_res += res_right
            _boxes[i]['layout'] = 'single'
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    #返回排序和布局后的结果。
    return new_res



def add_section_layout(doc, layout_flag):
    """添加文档部分并设置布局格式"""
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '1' if layout_flag == 'single' else '2')

def add_image_to_doc(doc, path, flag):
    """将图像添加到文档中"""
    paragraph_pic = doc.add_paragraph()
    paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph_pic.add_run("")
    run.add_picture(path, width=shared.Inches(5 if flag == 1 else 2))
